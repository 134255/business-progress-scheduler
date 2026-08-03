from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "技术落地规划.md"
OUTPUT = ROOT / "outputs" / "019fc5c9" / "业务进度管理微信小程序技术落地规划.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)
INK = RGBColor(31, 41, 55)
TABLE_HEADER = "E8EEF5"
BORDER = "D7DEE8"


def set_run_font(run, size=11, bold=None, color=INK, name="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for side, value in (("top", 90), ("bottom", 90), ("start", 120), ("end", 120)):
        node = cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), BORDER)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = paragraph.add_run("第 ")
    set_run_font(label, 9, color=MUTED)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, color=MUTED)


def add_inline(paragraph, text, size=11, color=INK):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=DARK_BLUE, name="Consolas", east_asia="Microsoft YaHei")
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def apply_style_tokens(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_masthead(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("业务进度管理微信小程序")
    set_run_font(run, size=24, bold=True, color=RGBColor(0, 0, 0))

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("技术落地规划与 MVP 实施基线")
    set_run_font(run, size=14, color=MUTED)

    metadata = [
        ("版本", "V0.1"),
        ("日期", "2026-08-03"),
        ("目标", "8 周内完成体验版与试运行准备"),
        ("技术基线", "原生微信小程序 + CloudBase 云开发"),
        ("状态", "开发已启动"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, 10.5, bold=True, color=INK)
        r2 = p.add_run(value)
        set_run_font(r2, 10.5, color=INK)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(16)
    note.paragraph_format.space_after = Pt(16)
    note.paragraph_format.left_indent = Inches(0.12)
    note.paragraph_format.right_indent = Inches(0.12)
    add_inline(note, "推荐结论：以顺序节点工作流作为 MVP，关键写操作全部经过云函数，凭证存云存储，通知采用订阅消息 + 站内待办双通道。", size=10.5, color=DARK_BLUE)


def parse_markdown(doc, source_text):
    lines = source_text.splitlines()
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        line = raw.strip()
        if not line or line.startswith("# 业务进度管理微信小程序技术落地规划") or line.startswith("版本：") or line.startswith("规划日期：") or line.startswith("目标："):
            index += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:], size=13, color=BLUE)
            index += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:], size=16, color=BLUE)
            index += 1
            continue

        if line.startswith("| ") and index + 1 < len(lines) and set(lines[index + 1].replace("|", "").replace("-", "").replace(":", "").strip()) == set():
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", cell) for cell in cells):
                    rows.append(cells)
                index += 1
            if rows:
                cols = len(rows[0])
                table = doc.add_table(rows=len(rows), cols=cols)
                if cols == 4:
                    widths = [3840, 1840, 1840, 1840]
                elif cols == 3:
                    widths = [4320, 2520, 2520]
                else:
                    widths = [9360 // cols] * cols
                    widths[-1] += 9360 - sum(widths)
                set_table_geometry(table, widths)
                for r_idx, row in enumerate(rows):
                    for c_idx, value in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        if r_idx == 0:
                            set_cell_shading(cell, TABLE_HEADER)
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.15
                        add_inline(p, value, size=9.5, color=INK)
                        for run in p.runs:
                            if r_idx == 0:
                                run.bold = True
                    if r_idx == 0:
                        set_repeat_table_header(table.rows[0])
                after = doc.add_paragraph()
                after.paragraph_format.space_after = Pt(2)
            continue

        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", line))
            index += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            index += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        add_inline(p, line.replace("  ", ""))
        index += 1


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    apply_style_tokens(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("业务进度管理微信小程序  |  技术落地规划")
    set_run_font(run, 9, color=MUTED)
    add_page_field(section.footer.paragraphs[0])

    add_masthead(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))

    core = doc.core_properties
    core.title = "业务进度管理微信小程序技术落地规划"
    core.subject = "微信小程序 MVP 技术方案、权限、数据、测试与发布规划"
    core.author = "项目组"
    core.keywords = "微信小程序, CloudBase, 业务进度, 节点管理, 凭证"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

